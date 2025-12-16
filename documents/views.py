# documents/views.py
from django.shortcuts import render, get_object_or_404
from django.http import HttpResponse
from django.contrib.auth.decorators import login_required
from io import BytesIO
import zipfile
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

from academics.models import Cohort, CourseSession
from students.models import Student, Enrollment
from finance.models import Installment


@login_required
def select_cohort(request):
    """Vue pour sélectionner une cohorte et générer les documents"""
    cohorts = Cohort.objects.all().order_by('-academic_year__start_date', 'name')

    context = {
        'cohorts': cohorts,
    }
    return render(request, 'documents/select_cohort.html', context)


@login_required
def generate_documents(request, cohort_id):
    """Génère et télécharge le ZIP contenant tous les documents"""
    cohort = get_object_or_404(Cohort, id=cohort_id)

    # Créer un buffer pour le ZIP en mémoire
    zip_buffer = BytesIO()

    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        # Créer le chemin de base
        base_path = f"{cohort.subject.name}/{cohort.name}"

        # === PACK ENSEIGNANT ===
        teacher_path = f"{base_path}/Pack_Enseignant"

        # 1. Fiche de Suivi des Séances
        suivi_doc = generate_suivi_seances(cohort)
        suivi_buffer = BytesIO()
        suivi_doc.save(suivi_buffer)
        suivi_buffer.seek(0)
        zip_file.writestr(f"{teacher_path}/Suivi_Seances.docx", suivi_buffer.read())

        # 2. Feuille d'Émargement
        emargement_doc = generate_emargement(cohort)
        emargement_buffer = BytesIO()
        emargement_doc.save(emargement_buffer)
        emargement_buffer.seek(0)
        zip_file.writestr(f"{teacher_path}/Liste_Presence.docx", emargement_buffer.read())

        # === PACK ÉTUDIANTS ===
        student_path = f"{base_path}/Pack_Etudiants"

        enrollments = Enrollment.objects.filter(cohort=cohort).select_related('student', 'tariff')

        for enrollment in enrollments:
            student_doc = generate_student_fiche(enrollment, cohort)
            student_buffer = BytesIO()
            student_doc.save(student_buffer)
            student_buffer.seek(0)

            student_name = f"{enrollment.student.last_name}_{enrollment.student.first_name}"
            zip_file.writestr(f"{student_path}/Fiche_{student_name}.docx", student_buffer.read())

    # Préparer la réponse HTTP
    zip_buffer.seek(0)
    response = HttpResponse(zip_buffer.read(), content_type='application/zip')
    filename = f"Documents_{cohort.subject.name}_{cohort.name}.zip"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'

    return response


def generate_suivi_seances(cohort):
    """Génère la Fiche de Suivi des Séances"""
    doc = Document()

    # En-tête
    heading = doc.add_heading(f'Fiche de Suivi des Séances', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Infos du groupe
    info_para = doc.add_paragraph()
    info_para.add_run(f"Matière: ").bold = True
    info_para.add_run(f"{cohort.subject.name}\n")
    info_para.add_run(f"Groupe: ").bold = True
    info_para.add_run(f"{cohort.name}\n")
    info_para.add_run(f"Professeur: ").bold = True
    info_para.add_run(f"{cohort.teacher.get_full_name()}\n")
    info_para.add_run(f"Période: ").bold = True
    info_para.add_run(f"{cohort.start_date.strftime('%d/%m/%Y')} - {cohort.end_date.strftime('%d/%m/%Y')}\n")

    doc.add_paragraph()

    # Récupérer les séances
    sessions = CourseSession.objects.filter(cohort=cohort).order_by('date', 'start_time')

    # Format plus spacieux : une section par séance
    for idx, session in enumerate(sessions, 1):
        # Ajouter un séparateur si ce n'est pas la première séance
        if idx > 1:
            doc.add_paragraph()

        # Créer un tableau pour chaque séance (plus d'espace)
        session_table = doc.add_table(rows=5, cols=2)
        session_table.style = 'Light Grid Accent 1'

        # Ligne 1: Numéro et Date
        session_table.rows[0].cells[0].text = f"Séance n°{idx}"
        session_table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
        session_table.rows[0].cells[1].text = session.date.strftime('%d/%m/%Y')

        # Ligne 2: Horaire
        session_table.rows[1].cells[0].text = "Horaire"
        session_table.rows[1].cells[0].paragraphs[0].runs[0].font.bold = True
        session_table.rows[1].cells[1].text = f"{session.start_time.strftime('%H:%M')} - {session.end_time.strftime('%H:%M')}"

        # Ligne 3: Contenu (avec beaucoup d'espace)
        content_cell = session_table.rows[2].cells[0]
        content_cell.text = "Contenu / Notes"
        content_cell.paragraphs[0].runs[0].font.bold = True

        # Fusionner les cellules pour le contenu sur toute la largeur
        content_merged = session_table.rows[2].cells[0].merge(session_table.rows[2].cells[1])

        # Zone de contenu avec espace pour écrire
        session_table.rows[3].cells[0].merge(session_table.rows[3].cells[1])
        content_area = session_table.rows[3].cells[0]

        # Ajouter le contenu existant ou laisser vide pour notes manuscrites
        if session.note:
            content_area.text = session.note
        else:
            # Ajouter plusieurs lignes vides pour permettre l'écriture manuscrite
            content_area.text = "\n\n\n"

        # Ajuster la hauteur de la cellule de contenu
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        tc = content_area._element
        tcPr = tc.get_or_add_tcPr()
        tcH = OxmlElement('w:tcH')
        tcH.set(qn('w:val'), '1500')  # Hauteur en twips (1500 = ~2.5cm)
        tcH.set(qn('w:hRule'), 'atLeast')
        tcPr.append(tcH)

        # Ligne 4: Signatures
        session_table.rows[4].cells[0].text = "Signature Professeur:"
        session_table.rows[4].cells[1].text = "Signature Administration:"

    return doc


def generate_emargement(cohort):
    """Génère la Feuille d'Émargement (présence)"""
    doc = Document()

    # Orientation paysage
    section = doc.sections[0]
    section.page_width = Inches(11.69)  # A4 paysage
    section.page_height = Inches(8.27)

    # En-tête
    heading = doc.add_heading(f"Feuille d'Émargement", 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Infos du groupe
    info_para = doc.add_paragraph()
    info_para.add_run(f"Groupe: ").bold = True
    info_para.add_run(f"{cohort.name} - {cohort.subject.name}\n")
    info_para.add_run(f"Professeur: ").bold = True
    info_para.add_run(f"{cohort.teacher.get_full_name()}\n")

    doc.add_paragraph()

    # Récupérer étudiants et séances
    enrollments = Enrollment.objects.filter(cohort=cohort).select_related('student').order_by('student__last_name')
    sessions = CourseSession.objects.filter(cohort=cohort).order_by('date', 'start_time')

    if not enrollments.exists():
        doc.add_paragraph("Aucun étudiant inscrit dans ce groupe.")
        return doc

    if not sessions.exists():
        doc.add_paragraph("Aucune séance prévue pour ce groupe.")
        return doc

    # Créer le tableau (Étudiants x Dates)
    num_students = enrollments.count()
    num_sessions = sessions.count()

    table = doc.add_table(rows=num_students + 1, cols=num_sessions + 2)
    table.style = 'Light Grid Accent 1'

    # En-tête: première colonne = N°, deuxième = Nom, suivantes = Dates
    header_row = table.rows[0]
    header_row.cells[0].text = "N°"
    header_row.cells[1].text = "Nom de l'étudiant"

    for idx, session in enumerate(sessions):
        cell = header_row.cells[idx + 2]
        cell.text = session.date.strftime('%d/%m')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8)
                run.font.bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Lignes des étudiants
    for student_idx, enrollment in enumerate(enrollments):
        row = table.rows[student_idx + 1]
        row.cells[0].text = str(student_idx + 1)
        row.cells[1].text = f"{enrollment.student.last_name} {enrollment.student.first_name}"

        # Cases vides pour signatures
        for session_idx in range(num_sessions):
            row.cells[session_idx + 2].text = ""

    return doc


def generate_student_fiche(enrollment, cohort):
    """Génère la Fiche d'Inscription / Financière pour un étudiant"""
    doc = Document()

    student = enrollment.student

    # En-tête
    heading = doc.add_heading("Fiche d'Inscription", 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # === INFORMATIONS ÉCOLE ===
    school_section = doc.add_heading("Informations de l'école", 2)
    school_para = doc.add_paragraph()
    school_para.add_run("École de Langues\n")
    school_para.add_run("Adresse: [À compléter]\n")
    school_para.add_run("Téléphone: [À compléter]\n")

    doc.add_paragraph()

    # === INFORMATIONS ÉTUDIANT ===
    student_section = doc.add_heading("Informations de l'étudiant", 2)
    student_table = doc.add_table(rows=6, cols=2)
    student_table.style = 'Light Grid Accent 1'

    student_data = [
        ("Nom", student.last_name),
        ("Prénom", student.first_name),
        ("Date de naissance", student.birth_date.strftime('%d/%m/%Y') if student.birth_date else "Non renseignée"),
        ("Téléphone", student.phone),
        ("Email", student.email or "Non renseigné"),
        ("Code étudiant", student.student_code),
    ]

    for idx, (label, value) in enumerate(student_data):
        row = student_table.rows[idx]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[1].text = value

    doc.add_paragraph()

    # === INFORMATIONS DU GROUPE ===
    group_section = doc.add_heading("Informations du cours", 2)
    group_table = doc.add_table(rows=6, cols=2)
    group_table.style = 'Light Grid Accent 1'

    # Construire le résumé des horaires à partir des WeeklySchedule
    weekly_schedules = cohort.weekly_schedules.all()
    if weekly_schedules.exists():
        schedule_parts = []
        for ws in weekly_schedules:
            day_name = ws.get_day_of_week_display()
            time_range = f"{ws.start_time.strftime('%H:%M')}-{ws.end_time.strftime('%H:%M')}"
            schedule_parts.append(f"{day_name} {time_range}")
        schedule = ", ".join(schedule_parts)
    else:
        schedule = "Non défini"

    group_data = [
        ("Matière", cohort.subject.name),
        ("Niveau", cohort.level.name if cohort.level else "Non spécifié"),
        ("Groupe", cohort.name),
        ("Professeur", cohort.teacher.get_full_name()),
        ("Horaires", schedule),
        ("Période", f"{cohort.start_date.strftime('%d/%m/%Y')} - {cohort.end_date.strftime('%d/%m/%Y')}"),
    ]

    for idx, (label, value) in enumerate(group_data):
        row = group_table.rows[idx]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[1].text = value

    doc.add_paragraph()

    # === INFORMATIONS FINANCIÈRES ===
    finance_section = doc.add_heading("Informations financières", 2)

    finance_para = doc.add_paragraph()
    finance_para.add_run("Tarif appliqué: ").bold = True
    finance_para.add_run(f"{enrollment.tariff.name} - {enrollment.tariff.amount} DA\n")

    finance_para.add_run("Mode de paiement: ").bold = True
    finance_para.add_run(f"{enrollment.get_payment_plan_display()}\n")

    finance_para.add_run("Montant total: ").bold = True
    finance_para.add_run(f"{enrollment.tariff.amount} DA\n")

    # Calculer le montant payé et le solde
    total_paid = sum(p.amount for p in enrollment.payments.all())
    balance = enrollment.balance_due

    finance_para.add_run("Montant payé: ").bold = True
    finance_para.add_run(f"{total_paid} DA\n")

    finance_para.add_run("Solde restant: ").bold = True
    if balance > 0:
        run = finance_para.add_run(f"{balance} DA\n")
        run.font.color.rgb = RGBColor(255, 0, 0)  # Rouge
        run.bold = True
    else:
        run = finance_para.add_run("SOLDÉ\n")
        run.font.color.rgb = RGBColor(0, 128, 0)  # Vert
        run.bold = True

    doc.add_paragraph()

    # === ÉCHÉANCIER DE PAIEMENT ===
    if enrollment.payment_plan == 'MONTHLY':
        payment_section = doc.add_heading("Échéancier de paiement", 3)

        installments = Installment.objects.filter(enrollment=enrollment).order_by('due_date')

        if installments.exists():
            installment_table = doc.add_table(rows=installments.count() + 1, cols=4)
            installment_table.style = 'Light Grid Accent 1'

            # En-tête
            header_row = installment_table.rows[0]
            header_row.cells[0].text = "Échéance"
            header_row.cells[1].text = "Date"
            header_row.cells[2].text = "Montant"
            header_row.cells[3].text = "Statut"

            for cell in header_row.cells:
                cell.paragraphs[0].runs[0].font.bold = True

            # Lignes des échéances
            for idx, installment in enumerate(installments):
                row = installment_table.rows[idx + 1]
                row.cells[0].text = f"Mois {idx + 1}"
                row.cells[1].text = installment.due_date.strftime('%d/%m/%Y')
                row.cells[2].text = f"{installment.amount} DA"

                if installment.is_paid:
                    status_para = row.cells[3].paragraphs[0]
                    run = status_para.add_run("✓ RÉGLÉ")
                    run.font.color.rgb = RGBColor(0, 128, 0)
                    run.bold = True
                else:
                    row.cells[3].text = "En attente"

    elif enrollment.payment_plan == 'FULL':
        if balance == 0:
            full_para = doc.add_paragraph()
            run = full_para.add_run("✓ Totalité réglée")
            run.font.color.rgb = RGBColor(0, 128, 0)
            run.bold = True
            run.font.size = Pt(14)

    # === HISTORIQUE DES PAIEMENTS ===
    payments = enrollment.payments.all().order_by('date')
    if payments.exists():
        doc.add_paragraph()
        payment_history_section = doc.add_heading("Historique des paiements", 3)

        payment_table = doc.add_table(rows=payments.count() + 1, cols=3)
        payment_table.style = 'Light Grid Accent 1'

        # En-tête
        header_row = payment_table.rows[0]
        header_row.cells[0].text = "Date"
        header_row.cells[1].text = "Montant"
        header_row.cells[2].text = "Méthode"

        for cell in header_row.cells:
            cell.paragraphs[0].runs[0].font.bold = True

        # Lignes des paiements
        for idx, payment in enumerate(payments):
            row = payment_table.rows[idx + 1]
            row.cells[0].text = payment.date.strftime('%d/%m/%Y')
            row.cells[1].text = f"{payment.amount} DA"
            row.cells[2].text = payment.get_method_display()

    # === SIGNATURES ===
    doc.add_paragraph()
    doc.add_paragraph()

    signature_table = doc.add_table(rows=2, cols=2)
    signature_table.rows[0].cells[0].text = "Signature de l'étudiant:"
    signature_table.rows[0].cells[1].text = "Signature de l'administration:"
    signature_table.rows[1].cells[0].text = "\n\n"
    signature_table.rows[1].cells[1].text = "\n\n"

    return doc
